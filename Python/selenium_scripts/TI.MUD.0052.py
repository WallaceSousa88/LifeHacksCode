import pyautogui
import time
import sys
import os
import json

from datetime import datetime, timedelta
from docx import Document
from docx.shared import Inches
from selenium import webdriver
from selenium.webdriver.edge.service import Service
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.common.keys import Keys

import json
with open('config.json', 'r') as f:
    config = json.load(f)

url_base = config['url_base']

user_id = sys.argv[1]
unidade_disco = sys.argv[2]

driver_path = f"{unidade_disco}:\\edgedriver_win64\\msedgedriver.exe"
driver = webdriver.Edge(service=Service(executable_path=driver_path))
driver.maximize_window()

screenshot_dir = f"{unidade_disco}:\\Users\\{user_id}\\Desktop\\Relatorio"
doc = Document(f"{unidade_disco}:\\Users\\{user_id}\\Desktop\\Relatorio\\TI.MUD.0052.docx")

driver.get(f"{url_base}GPJuri/_settings/repositories")
time.sleep(15)
WebDriverWait(driver, 20).until(EC.element_to_be_clickable((By.XPATH, "//a[@role='row' and @aria-rowindex='1']"))).click()
time.sleep(5)
WebDriverWait(driver, 20).until(EC.element_to_be_clickable((By.XPATH, "//span[text()='Policies']"))).click()
time.sleep(5)
WebDriverWait(driver, 20).until(EC.element_to_be_clickable((By.XPATH, "//span[@class='text-ellipsis' and text()='master']"))).click()
time.sleep(5)

screenshot1 = pyautogui.screenshot()
screenshot_name1 = "TI.MUD.0052_print_1.png"
screenshot_path1 = os.path.join(screenshot_dir, screenshot_name1)
screenshot1.save(screenshot_path1)
time.sleep(5)

element = driver.find_element(By.XPATH, "//div[@aria-level='1' and @class='bolt-header-title title-m l' and @role='heading' and text()='master']")
element.click()
time.sleep(5)

pyautogui.press('pagedown')
time.sleep(5)

screenshot2 = pyautogui.screenshot()
screenshot_name2 = "TI.MUD.0052_print_2.png"
screenshot_path2 = os.path.join(screenshot_dir, screenshot_name2)
screenshot2.save(screenshot_path2)
time.sleep(5)

driver.get(f"{url_base}_settings/audit")
time.sleep(5)

screenshot3 = pyautogui.screenshot()
screenshot_name3 = "TI.MUD.0052_print_3.png"
screenshot_path3 = os.path.join(screenshot_dir, screenshot_name3)
screenshot3.save(screenshot_path3)
time.sleep(5)

hoje = datetime.now()
ninety_days_ago = hoje - timedelta(days=90)

ninety_days_ago = ninety_days_ago.replace(hour=datetime.now().hour, minute=0) + timedelta(hours=4)
hoje = hoje.replace(hour=23, minute=59)
url = f"{url_base}_settings/audit?logs-period={ninety_days_ago.strftime('%Y-%m-%dT%H:%M')}Z-{hoje.strftime('%Y-%m-%dT%H:%M')}Z"
driver.get(url)
time.sleep(5)

csv = f"https://auditservice.dev.azure.com/cemig/_apis/audit/downloadLog?format=csv"
driver.get(csv)
time.sleep(5)

json = f"https://auditservice.dev.azure.com/cemig/_apis/audit/downloadLog?format=json"
driver.get(json)
time.sleep(5)

first_element_xpath = "//span[@class='text-ellipsis']"
first_element = WebDriverWait(driver, 20).until(
    EC.visibility_of_element_located((By.XPATH, first_element_xpath))
)
first_element.click()

scroll_time = 120
pause_time = 0.5
scrolls = int(scroll_time / pause_time)
for _ in range(scrolls):
    driver.find_element(By.TAG_NAME, 'body').send_keys(Keys.PAGE_DOWN)
    WebDriverWait(driver, pause_time).until(EC.visibility_of_element_located((By.TAG_NAME, 'body')))

screenshot4 = pyautogui.screenshot()
screenshot_name4 = "TI.MUD.0052_print_4.png"
screenshot_path4 = os.path.join(screenshot_dir, screenshot_name4)
screenshot4.save(screenshot_path4)
time.sleep(5)

def subs_texto_imagem(doc, search_text, image_path):
    for p in doc.paragraphs:
        if search_text in p.text:
            p.text = ''
            run = p.add_run()
            run.add_picture(image_path, width=Inches(6.0))

subs_texto_imagem(doc, '$04$', screenshot_path1)
subs_texto_imagem(doc, '$05$', screenshot_path2)
subs_texto_imagem(doc, '$06$', screenshot_path3)
subs_texto_imagem(doc, '$07$', screenshot_path4)

def subs_texto_data(doc, old_text, new_text):
    for p in doc.paragraphs:
        if old_text in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if old_text in inline[i].text:
                    text = inline[i].text.replace(old_text, new_text)
                    inline[i].text = text
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                subs_texto_data(cell, old_text, new_text)

now_str = datetime.now().strftime('%d/%m/%Y')
subs_texto_data(doc, '$01$', f'{hoje.strftime("%d/%m/%Y")}')
subs_texto_data(doc, '$02$', f'{ninety_days_ago.strftime("%d/%m/%Y")} a {hoje.strftime("%d/%m/%Y")}')
subs_texto_data(doc, '$03$', f'{ninety_days_ago.strftime("%d/%m/%Y")} a {hoje.strftime("%d/%m/%Y")}')

word_file_path = os.path.join(screenshot_dir, 'TI.MUD.0052.docx')
doc.save(word_file_path)

driver.quit()