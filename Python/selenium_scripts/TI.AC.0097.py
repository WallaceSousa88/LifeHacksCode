import pyautogui
import time
import sys
import os
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Inches
from selenium import webdriver
from selenium.webdriver.edge.service import Service
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.common.keys import Keys

user_id = sys.argv[1]
unidade_disco = sys.argv[2]

driver_path = f"{unidade_disco}:\\edgedriver_win64\\msedgedriver.exe"
driver = webdriver.Edge(service=Service(executable_path=driver_path))
driver.maximize_window()

screenshot_dir = f"{unidade_disco}:\\Users\\{user_id}\\Desktop\\Relatorio"
doc = Document(f"{unidade_disco}:\\Users\\{user_id}\\Desktop\\Relatorio\\TI.AC.0097.docx")

driver.get("https://dev.azure.com/cemig/_settings/organizationOverview")
time.sleep(15)

screenshot1 = pyautogui.screenshot()
screenshot_name1 = "TI.AC.0097_print_1.png"
screenshot_path1 = os.path.join(screenshot_dir, screenshot_name1)
screenshot1.save(screenshot_path1)
time.sleep(5)

driver.get("https://dev.azure.com/cemig/_settings/groups")
time.sleep(5)

screenshot2 = pyautogui.screenshot()
screenshot_name2 = "TI.AC.0097_print_2.png"
screenshot_path2 = os.path.join(screenshot_dir, screenshot_name2)
screenshot2.save(screenshot_path2)
time.sleep(5)

element = driver.find_element(By.XPATH, "//span[text()='Total']")
element.click()
time.sleep(5)

pyautogui.press('pagedown')
time.sleep(5)

screenshot3 = pyautogui.screenshot()
screenshot_name3 = "TI.AC.0097_print_3.png"
screenshot_path3 = os.path.join(screenshot_dir, screenshot_name3)
screenshot3.save(screenshot_path3)
time.sleep(5)   

search_field = WebDriverWait(driver, 20).until(EC.element_to_be_clickable((By.XPATH, "//input[@type='text' and @aria-label='Search groups']")))
search_field.click()
search_field.send_keys('Gerentes')
time.sleep(5)
cemig_element = WebDriverWait(driver, 20).until(EC.element_to_be_clickable((By.XPATH, "//span[@class='text-ellipsis capitalize' and text()='[cemig]']")))
cemig_element.click()
time.sleep(5)

screenshot4 = pyautogui.screenshot()
screenshot_name4 = "TI.AC.0097_print_4.png"
screenshot_path4 = os.path.join(screenshot_dir, screenshot_name4)
screenshot4.save(screenshot_path4)
time.sleep(5)

element = WebDriverWait(driver, 20).until(EC.element_to_be_clickable((By.XPATH, "//div[@class='text-ellipsis bolt-header-title title-m l' and @role='heading' and text()='[cemig]\\Gerentes']")))
element.click()
time.sleep(5)

pyautogui.press('pagedown')
time.sleep(5)

screenshot5 = pyautogui.screenshot()
screenshot_name5 = "TI.AC.0097_print_5.png"
screenshot_path5 = os.path.join(screenshot_dir, screenshot_name5)
screenshot5.save(screenshot_path5)
time.sleep(5)

driver.get("https://dev.azure.com/cemig/_settings/groups")
time.sleep(5)

search_field = WebDriverWait(driver, 20).until(EC.element_to_be_clickable((By.XPATH, "//input[@type='text' and @aria-label='Search groups']")))
search_field.click()
search_field.send_keys('NSIG - Coordenadores')
time.sleep(5)
cemig_element = WebDriverWait(driver, 20).until(EC.element_to_be_clickable((By.XPATH, "//span[@class='text-ellipsis capitalize' and text()='[cemig]']")))
cemig_element.click()
time.sleep(5)

screenshot6 = pyautogui.screenshot()
screenshot_name6 = "TI.AC.0097_print_6.png"
screenshot_path6 = os.path.join(screenshot_dir, screenshot_name6)
screenshot6.save(screenshot_path6)
time.sleep(5)

element = WebDriverWait(driver, 20).until(EC.element_to_be_clickable((By.XPATH, "//div[@class='text-ellipsis bolt-header-title title-m l' and @role='heading' and text()='[cemig]\\NSIG - Coordenadores']")))
element.click()
time.sleep(5)

pyautogui.press('pagedown')
time.sleep(5)

screenshot7 = pyautogui.screenshot()
screenshot_name7 = "TI.AC.0097_print_7.png"
screenshot_path7 = os.path.join(screenshot_dir, screenshot_name7)
screenshot7.save(screenshot_path7)
time.sleep(5)

driver.get("https://dev.azure.com/cemig/_settings/groups")
time.sleep(5)

search_field = WebDriverWait(driver, 20).until(EC.element_to_be_clickable((By.XPATH, "//input[@type='text' and @aria-label='Search groups']")))
search_field.click()
search_field.send_keys('Project Collection Administrators')
time.sleep(5)
cemig_element = WebDriverWait(driver, 20).until(EC.element_to_be_clickable((By.XPATH, "//span[@class='text-ellipsis capitalize' and text()='[cemig]']")))
cemig_element.click()
time.sleep(5)

screenshot8 = pyautogui.screenshot()
screenshot_name8 = "TI.AC.0097_print_8.png"
screenshot_path8 = os.path.join(screenshot_dir, screenshot_name8)
screenshot8.save(screenshot_path8)
time.sleep(5)

element = WebDriverWait(driver, 20).until(EC.element_to_be_clickable((By.XPATH, "//div[@class='text-ellipsis bolt-header-title title-m l' and @role='heading' and text()='[cemig]\\Project Collection Administrators']")))
element.click()
time.sleep(5)

pyautogui.press('pagedown')
time.sleep(5)

screenshot9 = pyautogui.screenshot()
screenshot_name9 = "TI.AC.0097_print_9.png"
screenshot_path9 = os.path.join(screenshot_dir, screenshot_name9)
screenshot9.save(screenshot_path9)
time.sleep(5)

WebDriverWait(driver, 20).until(EC.element_to_be_clickable((By.XPATH, "//span[text()='Members']"))).click()
time.sleep(5)

screenshot10 = pyautogui.screenshot()
screenshot_name10 = "TI.AC.0097_print_10.png"
screenshot_path10 = os.path.join(screenshot_dir, screenshot_name10)
screenshot10.save(screenshot_path10)
time.sleep(5)

element = WebDriverWait(driver, 20).until(EC.element_to_be_clickable((By.XPATH, "//div[@class='text-ellipsis bolt-header-title title-m l' and @role='heading' and text()='[cemig]\\Project Collection Administrators']")))
element.click()
time.sleep(5)

pyautogui.press('pagedown')
time.sleep(5)

screenshot11 = pyautogui.screenshot()
screenshot_name11 = "TI.AC.0097_print_11.png"
screenshot_path11 = os.path.join(screenshot_dir, screenshot_name11)
screenshot11.save(screenshot_path11)
time.sleep(5)

driver.get("https://dev.azure.com/cemig/_settings/groups")
time.sleep(5)

search_field = WebDriverWait(driver, 20).until(EC.element_to_be_clickable((By.XPATH, "//input[@type='text' and @aria-label='Search groups']")))
search_field.click()
search_field.send_keys('Project Collection Build Administrators')
time.sleep(5)
cemig_element = WebDriverWait(driver, 20).until(EC.element_to_be_clickable((By.XPATH, "//span[@class='text-ellipsis capitalize' and text()='[cemig]']")))
cemig_element.click()
time.sleep(5)

screenshot12 = pyautogui.screenshot()
screenshot_name12 = "TI.AC.0097_print_12.png"
screenshot_path12 = os.path.join(screenshot_dir, screenshot_name12)
screenshot12.save(screenshot_path12)
time.sleep(5)

element = WebDriverWait(driver, 20).until(EC.element_to_be_clickable((By.XPATH, "//div[@class='text-ellipsis bolt-header-title title-m l' and @role='heading' and text()='[cemig]\\Project Collection Build Administrators']")))
element.click()
time.sleep(5)

pyautogui.press('pagedown')
time.sleep(5)

screenshot13 = pyautogui.screenshot()
screenshot_name13 = "TI.AC.0097_print_13.png"
screenshot_path13 = os.path.join(screenshot_dir, screenshot_name13)
screenshot13.save(screenshot_path13)
time.sleep(5)

driver.get("https://dev.azure.com/cemig/_settings/groups")
time.sleep(5)

search_field = WebDriverWait(driver, 20).until(EC.element_to_be_clickable((By.XPATH, "//input[@type='text' and @aria-label='Search groups']")))
search_field.click()
search_field.send_keys('Project Collection Build Service Accounts')
time.sleep(5)
cemig_element = WebDriverWait(driver, 20).until(EC.element_to_be_clickable((By.XPATH, "//span[@class='text-ellipsis capitalize' and text()='[cemig]']")))
cemig_element.click()
time.sleep(5)

screenshot14 = pyautogui.screenshot()
screenshot_name14 = "TI.AC.0097_print_14.png"
screenshot_path14 = os.path.join(screenshot_dir, screenshot_name14)
screenshot14.save(screenshot_path14)
time.sleep(5)

element = WebDriverWait(driver, 20).until(EC.element_to_be_clickable((By.XPATH, "//div[@class='text-ellipsis bolt-header-title title-m l' and @role='heading' and text()='[cemig]\\Project Collection Build Service Accounts']")))
element.click()
time.sleep(5)

pyautogui.press('pagedown')
time.sleep(5)

screenshot15 = pyautogui.screenshot()
screenshot_name15 = "TI.AC.0097_print_15.png"
screenshot_path15 = os.path.join(screenshot_dir, screenshot_name15)
screenshot15.save(screenshot_path15)
time.sleep(5)

driver.get("https://dev.azure.com/cemig/_settings/groups")
time.sleep(5)

search_field = WebDriverWait(driver, 20).until(EC.element_to_be_clickable((By.XPATH, "//input[@type='text' and @aria-label='Search groups']")))
search_field.click()
search_field.send_keys('Project Collection Proxy Service Accounts')
time.sleep(5)
cemig_element = WebDriverWait(driver, 20).until(EC.element_to_be_clickable((By.XPATH, "//span[@class='text-ellipsis capitalize' and text()='[cemig]']")))
cemig_element.click()
time.sleep(5)

screenshot16 = pyautogui.screenshot()
screenshot_name16 = "TI.AC.0097_print_16.png"
screenshot_path16 = os.path.join(screenshot_dir, screenshot_name16)
screenshot16.save(screenshot_path16)
time.sleep(5)

element = WebDriverWait(driver, 20).until(EC.element_to_be_clickable((By.XPATH, "//div[@class='text-ellipsis bolt-header-title title-m l' and @role='heading' and text()='[cemig]\\Project Collection Proxy Service Accounts']")))
element.click()
time.sleep(5)

pyautogui.press('pagedown')
time.sleep(5)

screenshot17 = pyautogui.screenshot()
screenshot_name17 = "TI.AC.0097_print_17.png"
screenshot_path17 = os.path.join(screenshot_dir, screenshot_name17)
screenshot17.save(screenshot_path17)
time.sleep(5)

driver.get("https://dev.azure.com/cemig/_settings/groups")
time.sleep(5)

search_field = WebDriverWait(driver, 20).until(EC.element_to_be_clickable((By.XPATH, "//input[@type='text' and @aria-label='Search groups']")))
search_field.click()
search_field.send_keys('Project Collection Service Accounts')
time.sleep(5)
cemig_element = WebDriverWait(driver, 20).until(EC.element_to_be_clickable((By.XPATH, "//span[@class='text-ellipsis capitalize' and text()='[cemig]']")))
cemig_element.click()
time.sleep(5)

screenshot18 = pyautogui.screenshot()
screenshot_name18 = "TI.AC.0097_print_18.png"
screenshot_path18 = os.path.join(screenshot_dir, screenshot_name18)
screenshot18.save(screenshot_path18)
time.sleep(5)

element = WebDriverWait(driver, 20).until(EC.element_to_be_clickable((By.XPATH, "//div[@class='text-ellipsis bolt-header-title title-m l' and @role='heading' and text()='[cemig]\\Project Collection Service Accounts']")))
element.click()
time.sleep(5)

pyautogui.press('pagedown')
time.sleep(5)

screenshot19 = pyautogui.screenshot()
screenshot_name19 = "TI.AC.0097_print_19.png"
screenshot_path19 = os.path.join(screenshot_dir, screenshot_name19)
screenshot19.save(screenshot_path19)
time.sleep(5)

WebDriverWait(driver, 20).until(EC.element_to_be_clickable((By.XPATH, "//span[text()='Members']"))).click()
time.sleep(5)

screenshot20 = pyautogui.screenshot()
screenshot_name20 = "TI.AC.0097_print_20.png"
screenshot_path20 = os.path.join(screenshot_dir, screenshot_name20)
screenshot20.save(screenshot_path20)
time.sleep(5)

search_field = WebDriverWait(driver, 20).until(EC.element_to_be_clickable((By.XPATH, "//input[@type='text' and @aria-label='Search users and groups']")))
search_field.click()
search_field.send_keys('Enterprise Service Accounts')
time.sleep(5)
cemig_element = WebDriverWait(driver, 20).until(EC.element_to_be_clickable((By.XPATH, "//span[@class='text-ellipsis' and text()='group']")))
cemig_element.click()
time.sleep(5)

WebDriverWait(driver, 20).until(EC.element_to_be_clickable((By.XPATH, "//span[text()='Members']"))).click()
time.sleep(5)

screenshot21 = pyautogui.screenshot()
screenshot_name21 = "TI.AC.0097_print_21.png"
screenshot_path21 = os.path.join(screenshot_dir, screenshot_name21)
screenshot21.save(screenshot_path21)
time.sleep(5)

driver.get("https://dev.azure.com/cemig/_settings/groups")
time.sleep(5)

search_field = WebDriverWait(driver, 20).until(EC.element_to_be_clickable((By.XPATH, "//input[@type='text' and @aria-label='Search groups']")))
search_field.click()
search_field.send_keys('Project Collection Test Service Accounts')
time.sleep(5)
cemig_element = WebDriverWait(driver, 20).until(EC.element_to_be_clickable((By.XPATH, "//span[@class='text-ellipsis capitalize' and text()='[cemig]']")))
cemig_element.click()
time.sleep(5)

screenshot22 = pyautogui.screenshot()
screenshot_name22 = "TI.AC.0097_print_22.png"
screenshot_path22 = os.path.join(screenshot_dir, screenshot_name22)
screenshot22.save(screenshot_path22)
time.sleep(5)

element = WebDriverWait(driver, 20).until(EC.element_to_be_clickable((By.XPATH, "//div[@class='text-ellipsis bolt-header-title title-m l' and @role='heading' and text()='[cemig]\\Project Collection Test Service Accounts']")))
element.click()
time.sleep(5)

pyautogui.press('pagedown')
time.sleep(5)

screenshot23 = pyautogui.screenshot()
screenshot_name23 = "TI.AC.0097_print_23.png"
screenshot_path23 = os.path.join(screenshot_dir, screenshot_name23)
screenshot23.save(screenshot_path23)
time.sleep(5)

driver.get("https://dev.azure.com/cemig/_settings/groups")
time.sleep(5)

search_field = WebDriverWait(driver, 20).until(EC.element_to_be_clickable((By.XPATH, "//input[@type='text' and @aria-label='Search groups']")))
search_field.click()
search_field.send_keys('Project Collection Valid Users')
time.sleep(5)
cemig_element = WebDriverWait(driver, 20).until(EC.element_to_be_clickable((By.XPATH, "//span[@class='text-ellipsis capitalize' and text()='[cemig]']")))
cemig_element.click()
time.sleep(5)

screenshot24 = pyautogui.screenshot()
screenshot_name24 = "TI.AC.0097_print_24.png"
screenshot_path24 = os.path.join(screenshot_dir, screenshot_name24)
screenshot24.save(screenshot_path24)
time.sleep(5)

element = WebDriverWait(driver, 20).until(EC.element_to_be_clickable((By.XPATH, "//div[@class='text-ellipsis bolt-header-title title-m l' and @role='heading' and text()='[cemig]\\Project Collection Valid Users']")))
element.click()
time.sleep(5)

pyautogui.press('pagedown')
time.sleep(5)

screenshot25 = pyautogui.screenshot()
screenshot_name25 = "TI.AC.0097_print_25.png"
screenshot_path25 = os.path.join(screenshot_dir, screenshot_name25)
screenshot25.save(screenshot_path25)
time.sleep(5)

driver.get("https://dev.azure.com/cemig/_settings/groups")
time.sleep(5)

search_field = WebDriverWait(driver, 20).until(EC.element_to_be_clickable((By.XPATH, "//input[@type='text' and @aria-label='Search groups']")))
search_field.click()
search_field.send_keys('Project-Scoped Users')
time.sleep(5)
cemig_element = WebDriverWait(driver, 20).until(EC.element_to_be_clickable((By.XPATH, "//span[@class='text-ellipsis capitalize' and text()='[cemig]']")))
cemig_element.click()
time.sleep(5)

screenshot26 = pyautogui.screenshot()
screenshot_name26 = "TI.AC.0097_print_26.png"
screenshot_path26 = os.path.join(screenshot_dir, screenshot_name26)
screenshot26.save(screenshot_path26)
time.sleep(5)

element = WebDriverWait(driver, 20).until(EC.element_to_be_clickable((By.XPATH, "//div[@class='text-ellipsis bolt-header-title title-m l' and @role='heading' and text()='[cemig]\\Project-Scoped Users']")))
element.click()
time.sleep(5)

pyautogui.press('pagedown')
time.sleep(5)

screenshot27 = pyautogui.screenshot()
screenshot_name27 = "TI.AC.0097_print_27.png"
screenshot_path27 = os.path.join(screenshot_dir, screenshot_name27)
screenshot27.save(screenshot_path27)
time.sleep(5)

driver.get("https://dev.azure.com/cemig/_settings/groups")
time.sleep(5)

search_field = WebDriverWait(driver, 20).until(EC.element_to_be_clickable((By.XPATH, "//input[@type='text' and @aria-label='Search groups']")))
search_field.click()
search_field.send_keys('Qualidade')
time.sleep(5)
cemig_element = WebDriverWait(driver, 20).until(EC.element_to_be_clickable((By.XPATH, "//span[@class='text-ellipsis capitalize' and text()='[cemig]']")))
cemig_element.click()
time.sleep(5)

screenshot28 = pyautogui.screenshot()
screenshot_name28 = "TI.AC.0097_print_28.png"
screenshot_path28 = os.path.join(screenshot_dir, screenshot_name28)
screenshot28.save(screenshot_path28)
time.sleep(5)

element = WebDriverWait(driver, 20).until(EC.element_to_be_clickable((By.XPATH, "//div[@class='text-ellipsis bolt-header-title title-m l' and @role='heading' and text()='[cemig]\\Qualidade']")))
element.click()
time.sleep(5)

pyautogui.press('pagedown')
time.sleep(5)

screenshot29 = pyautogui.screenshot()
screenshot_name29 = "TI.AC.0097_print_29.png"
screenshot_path29 = os.path.join(screenshot_dir, screenshot_name29)
screenshot29.save(screenshot_path29)
time.sleep(5)

driver.get("https://dev.azure.com/cemig/_settings/groups")
time.sleep(5)

search_field = WebDriverWait(driver, 20).until(EC.element_to_be_clickable((By.XPATH, "//input[@type='text' and @aria-label='Search groups']")))
search_field.click()
search_field.send_keys('Security Service Group')
time.sleep(5)
cemig_element = WebDriverWait(driver, 20).until(EC.element_to_be_clickable((By.XPATH, "//span[@class='text-ellipsis capitalize' and text()='[cemig]']")))
cemig_element.click()
time.sleep(5)

screenshot30 = pyautogui.screenshot()
screenshot_name30 = "TI.AC.0097_print_30.png"
screenshot_path30 = os.path.join(screenshot_dir, screenshot_name30)
screenshot30.save(screenshot_path30)
time.sleep(5)

element = WebDriverWait(driver, 20).until(EC.element_to_be_clickable((By.XPATH, "//div[@class='text-ellipsis bolt-header-title title-m l' and @role='heading' and text()='[cemig]\\Security Service Group']")))
element.click()
time.sleep(5)

pyautogui.press('pagedown')
time.sleep(5)

screenshot31 = pyautogui.screenshot()
screenshot_name31 = "TI.AC.0097_print_31.png"
screenshot_path31 = os.path.join(screenshot_dir, screenshot_name31)
screenshot31.save(screenshot_path31)
time.sleep(5)

driver.get("https://dev.azure.com/cemig/GPJuri/_settings/permissions")
time.sleep(5)

screenshot32 = pyautogui.screenshot()
screenshot_name32 = "TI.AC.0097_print_32.png"
screenshot_path32 = os.path.join(screenshot_dir, screenshot_name32)
screenshot32.save(screenshot_path32)
time.sleep(5)

driver.get("https://dev.azure.com/cemig/GPJuri/_settings/repositories")
time.sleep(5)

element = WebDriverWait(driver, 20).until(EC.element_to_be_clickable((By.XPATH, "//span[@class='text-ellipsis flex-self-center' and text()='GPJuri']")))
element.click()
time.sleep(5)

element = WebDriverWait(driver, 20).until(EC.element_to_be_clickable((By.XPATH, "//span[@class='bolt-tab-text' and @data-content='Security']")))
element.click()
time.sleep(5)

element = WebDriverWait(driver, 20).until(EC.element_to_be_clickable((By.XPATH, "//span[@class='text-ellipsis' and text()='master']")))
element.click()
time.sleep(5)

screenshot33 = pyautogui.screenshot()
screenshot_name33 = "TI.AC.0097_print_33.png"
screenshot_path33 = os.path.join(screenshot_dir, screenshot_name33)
screenshot33.save(screenshot_path33)
time.sleep(5)

element = WebDriverWait(driver, 20).until(EC.element_to_be_clickable((By.XPATH, "//span[@class='capitalize members-text ' and text()='Build Administrators']")))
element.click()

screenshot34 = pyautogui.screenshot()
screenshot_name34 = "TI.AC.0097_print_34.png"
screenshot_path34 = os.path.join(screenshot_dir, screenshot_name34)
screenshot34.save(screenshot_path34)
time.sleep(5)

driver.get("https://dev.azure.com/cemig/GPJuri/_settings/repositories")
time.sleep(5)

element = WebDriverWait(driver, 20).until(EC.element_to_be_clickable((By.XPATH, "//span[@class='text-ellipsis flex-self-center' and text()='GPJuri']")))
element.click()
time.sleep(5)

element = WebDriverWait(driver, 20).until(EC.element_to_be_clickable((By.XPATH, "//span[@class='bolt-tab-text' and @data-content='Security']")))
element.click()
time.sleep(5)

element = WebDriverWait(driver, 20).until(EC.element_to_be_clickable((By.XPATH, "//span[@class='text-ellipsis' and text()='master']")))
element.click()
time.sleep(5)

element = WebDriverWait(driver, 20).until(EC.element_to_be_clickable((By.XPATH, "//span[@class='capitalize members-text ' and text()='Contributors']")))
element.click()

screenshot35 = pyautogui.screenshot()
screenshot_name35 = "TI.AC.0097_print_35.png"
screenshot_path35 = os.path.join(screenshot_dir, screenshot_name35)
screenshot35.save(screenshot_path35)
time.sleep(5)

driver.get("https://dev.azure.com/cemig/GPJuri/_settings/repositories")
time.sleep(5)

element = WebDriverWait(driver, 20).until(EC.element_to_be_clickable((By.XPATH, "//span[@class='text-ellipsis flex-self-center' and text()='GPJuri']")))
element.click()
time.sleep(5)

element = WebDriverWait(driver, 20).until(EC.element_to_be_clickable((By.XPATH, "//span[@class='bolt-tab-text' and @data-content='Security']")))
element.click()
time.sleep(5)

element = WebDriverWait(driver, 20).until(EC.element_to_be_clickable((By.XPATH, "//span[@class='text-ellipsis' and text()='master']")))
element.click()
time.sleep(5)

element = WebDriverWait(driver, 20).until(EC.element_to_be_clickable((By.XPATH, "//span[@class='capitalize members-text ' and text()='Project Administrators']")))
element.click()

screenshot36 = pyautogui.screenshot()
screenshot_name36 = "TI.AC.0097_print_36.png"
screenshot_path36 = os.path.join(screenshot_dir, screenshot_name36)
screenshot36.save(screenshot_path36)
time.sleep(5)

driver.get("https://dev.azure.com/cemig/GPJuri/_settings/projectOverview")
time.sleep(5)

screenshot37 = pyautogui.screenshot()
screenshot_name37 = "TI.AC.0097_print_37.png"
screenshot_path37 = os.path.join(screenshot_dir, screenshot_name37)
screenshot37.save(screenshot_path37)
time.sleep(5)

driver.get("https://dev.azure.com/cemig/GPJuri/_settings/repositories")
time.sleep(5)

element = WebDriverWait(driver, 20).until(EC.element_to_be_clickable((By.XPATH, "//span[@class='text-ellipsis flex-self-center' and text()='GPJuri']")))
element.click()
time.sleep(5)

element = WebDriverWait(driver, 20).until(EC.element_to_be_clickable((By.XPATH, "//span[@class='bolt-tab-text' and @data-content='Security']")))
element.click()
time.sleep(5)

element = WebDriverWait(driver, 20).until(EC.element_to_be_clickable((By.XPATH, "//span[@class='text-ellipsis' and text()='master']")))
element.click()
time.sleep(5)

element = WebDriverWait(driver, 20).until(EC.element_to_be_clickable((By.XPATH, "//span[@class='capitalize members-text ' and text()='Readers']")))
element.click()

screenshot38 = pyautogui.screenshot()
screenshot_name38 = "TI.AC.0097_print_38.png"
screenshot_path38 = os.path.join(screenshot_dir, screenshot_name38)
screenshot38.save(screenshot_path38)
time.sleep(5)

driver.get("https://dev.azure.com/cemig/GPJuri/_settings/repositories")
time.sleep(5)

element = WebDriverWait(driver, 20).until(EC.element_to_be_clickable((By.XPATH, "//span[@class='text-ellipsis flex-self-center' and text()='GPJuri']")))
element.click()
time.sleep(5)

element = WebDriverWait(driver, 20).until(EC.element_to_be_clickable((By.XPATH, "//span[@class='bolt-tab-text' and @data-content='Security']")))
element.click()
time.sleep(5)

element = WebDriverWait(driver, 20).until(EC.element_to_be_clickable((By.XPATH, "//span[@class='text-ellipsis' and text()='master']")))
element.click()
time.sleep(5)

element = WebDriverWait(driver, 20).until(EC.element_to_be_clickable((By.XPATH, "//span[@class='capitalize members-text ' and text()='FERNANDO LOPES DA SILVEIRA']")))
element.click()

screenshot39 = pyautogui.screenshot()
screenshot_name39 = "TI.AC.0097_print_39.png"
screenshot_path39 = os.path.join(screenshot_dir, screenshot_name39)
screenshot39.save(screenshot_path39)
time.sleep(5)

driver.get("https://dev.azure.com/cemig/_settings/groups")
time.sleep(5)

search_field = WebDriverWait(driver, 20).until(EC.element_to_be_clickable((By.XPATH, "//input[@type='text' and @aria-label='Search groups']")))
search_field.click()
search_field.send_keys('NSIG - Coordenadores')
time.sleep(5)
cemig_element = WebDriverWait(driver, 20).until(EC.element_to_be_clickable((By.XPATH, "//span[@class='text-ellipsis capitalize' and text()='[cemig]']")))
cemig_element.click()
time.sleep(5)

element = WebDriverWait(driver, 20).until(EC.element_to_be_clickable((By.XPATH, "//span[@class='bolt-tab-text' and @data-content='Members']")))
element.click()
time.sleep(5)

screenshot40 = pyautogui.screenshot()
screenshot_name40 = "TI.AC.0097_print_40.png"
screenshot_path40 = os.path.join(screenshot_dir, screenshot_name40)
screenshot40.save(screenshot_path40)
time.sleep(5)

driver.get("https://dev.azure.com/cemig/_settings/audit")
time.sleep(5)

screenshot41 = pyautogui.screenshot()
screenshot_name41 = "TI.MUD.0059_print_41.png"
screenshot_path41 = os.path.join(screenshot_dir, screenshot_name41)
screenshot41.save(screenshot_path41)
time.sleep(5)

hoje = datetime.now()
ninety_days_ago = hoje - timedelta(days=90)

ninety_days_ago = ninety_days_ago.replace(hour=datetime.now().hour, minute=0) + timedelta(hours=4)
hoje = hoje.replace(hour=23, minute=59)
url = f"https://dev.azure.com/cemig/_settings/audit?logs-period={ninety_days_ago.strftime('%Y-%m-%dT%H:%M')}Z-{hoje.strftime('%Y-%m-%dT%H:%M')}Z"
driver.get(url)
time.sleep(5)

csv = f"https://auditservice.dev.azure.com/cemig/_apis/audit/downloadLog?format=csv"
driver.get(csv)
time.sleep(5)

first_element_xpath = "//span[@class='text-ellipsis']"
first_element = WebDriverWait(driver, 20).until(
    EC.visibility_of_element_located((By.XPATH, first_element_xpath))
)
first_element.click()

scroll_time = 180
pause_time = 0.5
scrolls = int(scroll_time / pause_time)
for _ in range(scrolls):
    driver.find_element(By.TAG_NAME, 'body').send_keys(Keys.PAGE_DOWN)
    WebDriverWait(driver, pause_time).until(EC.visibility_of_element_located((By.TAG_NAME, 'body')))

screenshot42 = pyautogui.screenshot()
screenshot_name42 = "TI.AC.0097_print_42.png"
screenshot_path42 = os.path.join(screenshot_dir, screenshot_name42)
screenshot42.save(screenshot_path42)

def subs_texto_imagem(doc, search_text, image_path):
    for p in doc.paragraphs:
        if search_text in p.text:
            p.text = ''
            run = p.add_run()
            run.add_picture(image_path, width=Inches(6.0))

subs_texto_imagem(doc, '$03$', screenshot_path1)
subs_texto_imagem(doc, '$04$', screenshot_path2)
subs_texto_imagem(doc, '$05$', screenshot_path3)
subs_texto_imagem(doc, '$06$', screenshot_path4)
subs_texto_imagem(doc, '$07$', screenshot_path5)
subs_texto_imagem(doc, '$08$', screenshot_path6)
subs_texto_imagem(doc, '$09$', screenshot_path7)
subs_texto_imagem(doc, '$10$', screenshot_path8)
subs_texto_imagem(doc, '$11$', screenshot_path9)
subs_texto_imagem(doc, '$12$', screenshot_path10)
subs_texto_imagem(doc, '$13$', screenshot_path11)
subs_texto_imagem(doc, '$14$', screenshot_path12)
subs_texto_imagem(doc, '$15$', screenshot_path13)
subs_texto_imagem(doc, '$16$', screenshot_path14)
subs_texto_imagem(doc, '$17$', screenshot_path15)
subs_texto_imagem(doc, '$18$', screenshot_path16)
subs_texto_imagem(doc, '$19$', screenshot_path17)
subs_texto_imagem(doc, '$20$', screenshot_path18)
subs_texto_imagem(doc, '$21$', screenshot_path19)
subs_texto_imagem(doc, '$22$', screenshot_path20)
subs_texto_imagem(doc, '$23$', screenshot_path21)
subs_texto_imagem(doc, '$24$', screenshot_path22)
subs_texto_imagem(doc, '$25$', screenshot_path23)
subs_texto_imagem(doc, '$26$', screenshot_path24)
subs_texto_imagem(doc, '$27$', screenshot_path25)
subs_texto_imagem(doc, '$28$', screenshot_path26)
subs_texto_imagem(doc, '$29$', screenshot_path27)
subs_texto_imagem(doc, '$30$', screenshot_path28)
subs_texto_imagem(doc, '$31$', screenshot_path29)
subs_texto_imagem(doc, '$32$', screenshot_path30)
subs_texto_imagem(doc, '$33$', screenshot_path31)
subs_texto_imagem(doc, '$34$', screenshot_path32)
subs_texto_imagem(doc, '$35$', screenshot_path33)
subs_texto_imagem(doc, '$36$', screenshot_path34)
subs_texto_imagem(doc, '$37$', screenshot_path35)
subs_texto_imagem(doc, '$38$', screenshot_path36)
subs_texto_imagem(doc, '$39$', screenshot_path37)
subs_texto_imagem(doc, '$40$', screenshot_path38)
subs_texto_imagem(doc, '$41$', screenshot_path39)
subs_texto_imagem(doc, '$42$', screenshot_path40)

def subs_texto_data(doc, old_text, new_text):
    for p in doc.paragraphs:
        if old_text in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if old_text in inline[i].text:
                    text = inline[i].text.replace(old_text, new_text)
                    inline[i].text = text
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                subs_texto_data(cell, old_text, new_text)

now_str = datetime.now().strftime('%d/%m/%Y')
subs_texto_data(doc, '$01$', f'{hoje.strftime("%d/%m/%Y")}')
subs_texto_data(doc, '$02$', f'{ninety_days_ago.strftime("%d/%m/%Y")} a {hoje.strftime("%d/%m/%Y")}')

word_file_path = os.path.join(screenshot_dir, 'TI.AC.0097.docx')
doc.save(word_file_path)

driver.quit()
